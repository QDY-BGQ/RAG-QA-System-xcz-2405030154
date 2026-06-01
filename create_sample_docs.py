"""
创建示例文档脚本：生成5份关于自然语言处理的DOCX文档
"""

import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_document(title, content_list, filename):
    """创建DOCX文档"""
    doc = Document()
    
    heading = doc.add_heading(title, 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    for section_title, section_content in content_list:
        doc.add_heading(section_title, level=1)
        for para in section_content:
            doc.add_paragraph(para)
        doc.add_paragraph()
    
    doc.add_paragraph("---")
    doc.add_paragraph("本文档为RAG智能问答系统的示例测试文档。").italic = True
    
    output_path = os.path.join("docs", filename)
    doc.save(output_path)
    print(f"✓ 已创建: {output_path}")


def main():
    os.makedirs("docs", exist_ok=True)
    
    doc1_content = [
        ("一、什么是自然语言处理", [
            "自然语言处理（Natural Language Processing，简称NLP）是人工智能的一个重要分支，旨在使计算机能够理解、解释和生成人类语言。",
            "NLP结合了计算语言学、统计学和机器学习技术，用于处理和分析大量的自然语言数据。",
            "自然语言处理的最终目标是让计算机能够像人类一样理解和处理语言，实现人与计算机之间的自然交流。"
        ]),
        ("二、NLP的发展历程", [
            "自然语言处理的发展可以分为三个阶段：规则驱动阶段（1950s-1990s）、统计学习阶段（1990s-2010s）和深度学习阶段（2010s至今）。",
            "早期的NLP系统主要基于人工编写的规则，能够处理简单的语言任务，但难以应对复杂的语言现象。",
            "随着机器学习技术的发展，特别是深度学习的兴起，NLP取得了突破性进展，在机器翻译、文本分类、问答系统等领域达到或超过了人类水平。"
        ]),
        ("三、NLP的主要任务", [
            "自然语言处理包含多种任务，主要分为自然语言理解（NLU）和自然语言生成（NLG）两大类。",
            "常见的NLP任务包括：文本分类、情感分析、命名实体识别、关系抽取、机器翻译、文本摘要、问答系统等。",
            "这些任务广泛应用于搜索引擎、智能客服、机器翻译、内容审核、语音助手等领域。"
        ])
    ]
    
    doc2_content = [
        ("一、词嵌入的概念", [
            "词嵌入（Word Embedding）是将自然语言中的词汇映射为实数向量的技术，是深度学习时代NLP的基础。",
            "通过词嵌入，我们可以将离散的词汇转化为连续的向量空间中的点，使得语义相似的词汇在向量空间中距离更近。",
            "词嵌入的核心思想是'分布式表示'，即每个词的含义由其上下文共同决定。"
        ]),
        ("二、经典的词嵌入模型", [
            "Word2Vec是Google在2013年提出的词嵌入模型，包括CBOW（Continuous Bag-of-Words）和Skip-gram两种架构。",
            "GloVe（Global Vectors for Word Representation）是Stanford提出的基于全局词共现统计的词嵌入模型。",
            "ELMo是首个大规模预训练的上下文相关词嵌入模型，能够捕捉词汇的多义性。"
        ]),
        ("三、词嵌入的应用", [
            "词嵌入是几乎所有NLP深度学习模型的基础组件，为模型提供输入表示。",
            "通过词嵌入，我们可以计算词汇之间的语义相似度，进行词汇类比推理（如'国王-男人+女人≈王后'）。",
            "词嵌入还可以用于文本分类、信息检索、推荐系统等多种任务中。"
        ])
    ]
    
    doc3_content = [
        ("一、Transformer模型简介", [
            "Transformer是Google在2017年的论文《Attention Is All You Need》中提出的神经网络架构，完全基于自注意力机制。",
            "与传统的循环神经网络（RNN）不同，Transformer可以并行处理整个序列，大大提高了训练效率。",
            "Transformer已成为现代NLP的基础架构，GPT系列、BERT、T5等著名模型都基于Transformer。"
        ]),
        ("二、自注意力机制", [
            "自注意力机制（Self-Attention）是Transformer的核心，允许模型在处理每个位置时关注序列中的其他位置。",
            "通过计算查询（Query）、键（Key）和值（Value）的注意力权重，模型可以动态地分配对不同位置的关注度。",
            "多头注意力（Multi-Head Attention）通过多个独立的注意力头，可以捕捉不同子空间中的语义关系。"
        ]),
        ("三、Transformer的架构", [
            "Transformer由编码器（Encoder）和解码器（Decoder）两部分组成，编码器负责理解输入序列，解码器负责生成输出序列。",
            "位置编码（Positional Encoding）为模型提供序列的位置信息，因为Attention本身不包含顺序信息。",
            "残差连接和层归一化是Transformer中的重要组件，有助于训练深层网络。"
        ])
    ]
    
    doc4_content = [
        ("一、BERT模型简介", [
            "BERT（Bidirectional Encoder Representations from Transformers）是Google在2018年提出的预训练语言模型。",
            "BERT基于Transformer的编码器部分，通过双向预训练，能够深度理解上下文信息。",
            "BERT的出现标志着NLP进入了预训练模型时代，在11项NLP基准任务上取得了SOTA成绩。"
        ]),
        ("二、BERT的预训练任务", [
            "掩码语言模型（Masked Language Modeling，MLM）：随机遮盖序列中的部分词汇，让模型预测被遮盖的词。",
            "下一句预测（Next Sentence Prediction，NSP）：给定两个句子，预测它们是否是连续的上下文关系。",
            "这两个预训练任务使BERT能够同时学习词级和句子级的表示。"
        ]),
        ("三、BERT的微调与应用", [
            "BERT的微调非常简单，只需在预训练模型之上添加少量的任务特定层，然后用任务数据进行微调。",
            "BERT可以应用于几乎所有NLP任务，包括文本分类、命名实体识别、问答系统、自然语言推理等。",
            "BERT的变体包括RoBERTa、ALBERT、ELECTRA等，通过改进预训练策略或模型结构进一步提升性能。"
        ])
    ]
    
    doc5_content = [
        ("一、文本分类概述", [
            "文本分类是NLP中最基础也最重要的任务之一，目标是将文本自动分配到预定义的类别中。",
            "文本分类的应用场景包括：新闻分类、垃圾邮件检测、情感分析、用户评论分类、内容审核等。",
            "根据分类的粒度，文本分类可以分为二分类、多分类和多标签分类。"
        ]),
        ("二、传统的文本分类方法", [
            "基于规则的方法：通过人工编写的规则进行分类，如关键词匹配、正则表达式等。",
            "基于机器学习的方法：使用TF-IDF等特征提取方法，配合朴素贝叶斯、SVM、随机森林等分类器。",
            "传统方法的优点是可解释性强，但需要大量的特征工程，且难以捕捉深层语义。"
        ]),
        ("三、深度学习文本分类方法", [
            "CNN（卷积神经网络）：通过卷积层捕捉文本的n-gram特征，在文本分类任务上表现出色。",
            "RNN/LSTM：通过循环结构捕捉序列依赖关系，适合处理长文本。",
            "预训练模型（如BERT）：通过在大规模语料上预训练，然后微调，是当前文本分类的SOTA方法。"
        ])
    ]
    
    create_document("自然语言处理简介", doc1_content, "1_自然语言处理简介.docx")
    create_document("词嵌入技术详解", doc2_content, "2_词嵌入技术详解.docx")
    create_document("Transformer模型原理", doc3_content, "3_Transformer模型原理.docx")
    create_document("BERT模型详解", doc4_content, "4_BERT模型详解.docx")
    create_document("文本分类技术综述", doc5_content, "5_文本分类技术综述.docx")
    
    print("\n✓ 所有示例文档创建完成！")
    print(f"  共创建 {len(os.listdir('docs'))} 份文档")


if __name__ == "__main__":
    main()