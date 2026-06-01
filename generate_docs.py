from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

docs_content = [
    {
        'title': '自然语言处理技术概述',
        'filename': '1_自然语言处理概述.docx',
        'content': '''一、自然语言处理简介

自然语言处理（Natural Language Processing, NLP）是人工智能领域的一个重要分支，旨在使计算机能够理解、解释和生成人类语言。

二、NLP的主要任务

1. 文本分类
   - 将文本自动分类到预定义的类别中
   - 应用场景：垃圾邮件过滤、情感分析、主题分类

2. 命名实体识别（NER）
   - 识别文本中的实体，如人名、地名、机构名等
   - 应用场景：信息抽取、知识图谱构建

3. 语义分析
   - 理解句子的含义和上下文
   - 应用场景：问答系统、机器翻译

4. 文本生成
   - 自动生成自然语言文本
   - 应用场景：文章写作、对话系统

三、NLP发展历程

1. 规则驱动阶段（1950s-1980s）
   - 基于手工编写的语法规则
   - 代表性系统：ELIZA

2. 统计机器学习阶段（1990s-2010s）
   - 基于大规模语料库的统计方法
   - 代表性技术：隐马尔可夫模型、条件随机场

3. 深度学习阶段（2010s-至今）
   - 基于神经网络的端到端学习
   - 代表性模型：Word2Vec、Transformer、BERT

四、关键技术挑战

1. 歧义消解
   - 一词多义问题
   - 上下文理解

2. 上下文依赖
   - 长文本理解
   - 指代消解

3. 领域自适应
   - 跨领域知识迁移
   - 少样本学习

五、应用场景

- 智能客服系统
- 机器翻译
- 文本摘要
- 情感分析
- 信息检索
'''
    },
    {
        'title': '词嵌入技术详解',
        'filename': '2_词嵌入技术详解.docx',
        'content': '''一、词嵌入的概念

词嵌入（Word Embedding）是一种将离散词汇映射到连续向量空间的技术，使得语义相似的词在向量空间中距离较近。

二、传统词表示方法

1. 独热编码（One-hot Encoding）
   - 每个词对应一个维度
   - 维度爆炸问题
   - 无法表达语义关系

2. 词频-逆文档频率（TF-IDF）
   - 基于统计的权重计算
   - 忽略词序信息

三、经典词嵌入模型

1. Word2Vec（2013年，Mikolov等人）

   主要架构：
   - CBOW（Continuous Bag of Words）：根据上下文预测中心词
   - Skip-gram：根据中心词预测上下文

   训练目标：
   - 最大化对数似然概率
   - 使用负采样加速训练

2. GloVe（2014年，斯坦福大学）

   核心思想：
   - 基于全局词共现矩阵
   - 最小化加权平方损失
   - 结合了全局统计信息

3. FastText（2016年，Facebook）

   创新点：
   - 引入子词信息
   - 支持OOV词
   - 适合形态丰富的语言

四、词嵌入的特性

1. 语义相似性
   - 向量相似度反映语义相似度
   - 可进行向量运算

2. 类比推理
   - 展示词向量的结构特性

3. 上下文感知
   - 静态嵌入与动态嵌入

五、进阶发展

1. ELMo（2018年）
   - 基于双向LSTM的上下文相关嵌入

2. BERT嵌入
   - 基于Transformer的双向预训练

3. Sentence-BERT
   - 专门用于句子级嵌入

六、应用场景

- 文本分类
- 情感分析
- 机器翻译
- 问答系统
- 推荐系统
'''
    },
    {
        'title': 'Transformer模型原理',
        'filename': '3_Transformer模型原理.docx',
        'content': '''一、Transformer的诞生

Transformer是Google在2017年论文《Attention is All You Need》中提出的全新架构，彻底改变了NLP领域。

二、Transformer架构概览

整体结构：
- 编码器：6层堆叠
- 解码器：6层堆叠
- 输入嵌入层
- 位置编码层

三、核心组件：自注意力机制

1. 注意力机制原理

   注意力权重计算：
   Attention(Q, K, V) = softmax(QK^T / sqrt(d_k)) * V

   其中：
   - Q：查询向量
   - K：键向量
   - V：值向量

2. 多头注意力

   将Q、K、V线性变换h次，每个head学习不同的注意力模式。

3. 编码器中的注意力

   - 自注意力：输入序列与自身计算注意力

4. 解码器中的注意力

   - 掩码自注意力：防止未来信息泄露
   - 编码器-解码器注意力：关注编码器输出

四、位置编码

由于Transformer没有递归结构，需要显式编码位置信息。

特性：
- 固定的位置编码
- 相对位置关系可学习

五、前馈神经网络

每层包含两个线性变换和ReLU激活。

六、残差连接与层归一化

每个子层后添加残差连接。

七、Transformer的优势

1. 并行化能力
   - 相比RNN，可完全并行计算

2. 长距离依赖建模
   - 注意力机制直接建模任意位置关系

3. 统一框架
   - 编码器-解码器架构适用于多种任务

八、Transformer的变体

1. BERT：仅编码器架构
2. GPT：仅解码器架构
3. T5：统一文本到文本框架

九、应用影响

- 机器翻译质量大幅提升
- 预训练模型成为标准范式
- 推动NLP技术快速发展
'''
    },
    {
        'title': 'BERT模型详解',
        'filename': '4_BERT模型详解.docx',
        'content': '''一、BERT简介

BERT是Google在2018年提出的预训练模型，基于Transformer编码器架构。

二、BERT的核心创新

1. 双向上下文理解
   - 同时考虑左右上下文

2. Masked Language Model预训练任务
   - 随机掩盖15%的输入token
   - 预测被掩盖的token

3. Next Sentence Prediction任务
   - 判断两个句子是否连续

三、BERT架构

1. 模型尺寸

   - BERT-Base：12层Transformer，768维度，12头注意力
   - BERT-Large：24层Transformer，1024维度，16头注意力

2. 输入表示

   输入序列格式：<[BOS_never_used_51bce0c785ca2f68081bfa7d91973934]> token_1 token_2 ... token_n [SEP]

   嵌入层组合：Token Embeddings + Segment Embeddings + Position Embeddings

3. 输出

   - <[BOS_never_used_51bce0c785ca2f68081bfa7d91973934]> token用于分类任务
   - 其他token用于序列标注任务

四、预训练与微调流程

1. 预训练阶段
   - 使用大规模无标注文本
   - 同时优化MLM和NSP任务

2. 微调阶段
   - 在预训练模型基础上
   - 添加任务特定层

五、BERT的优势

1. 双向上下文
   - 优于单向模型

2. 通用表示
   - 单一模型适配多种任务

3. 迁移学习能力
   - 预训练知识可迁移

六、BERT的变体

1. ALBERT：参数共享减少模型大小
2. RoBERTa：移除NSP任务
3. DistilBERT：知识蒸馏版本
4. ELECTRA：判别式任务

七、BERT的影响

- 刷新多项NLP任务SOTA
- 开启预训练模型时代
- 成为NLP研究的基础模型
'''
    },
    {
        'title': '文本生成技术综述',
        'filename': '5_文本生成技术综述.docx',
        'content': '''一、文本生成概述

文本生成是NLP的核心任务之一，目标是让计算机自动生成符合人类语言习惯的文本内容。

二、传统文本生成方法

1. 基于规则的方法
   - 模板匹配
   - 语法规则

2. 统计机器翻译方法
   - n-gram模型
   - 隐马尔可夫模型

三、神经网络文本生成

1. 循环神经网络（RNN）
   - 序列建模能力
   - LSTM和GRU改进

2. Seq2Seq架构
   - 编码器-解码器结构
   - 注意力机制增强

3. Transformer解码器
   - 纯注意力架构
   - GPT系列模型

四、生成对抗网络（GAN）

1. GAN在文本生成中的应用
   - Generator生成文本
   - Discriminator判别真伪

2. SeqGAN
   - 序列级GAN
   - 强化学习训练

五、预训练语言模型

1. GPT系列
   - GPT-1到GPT-4的演进

2. T5
   - 统一文本到文本框架

3. BERT用于生成
   - BERT-GPT混合架构

六、文本生成评估指标

1. BLEU分数
2. ROUGE分数
3. METEOR
4. BERTScore

七、文本生成的挑战

1. 一致性问题
2. 多样性问题
3. 可控性问题
4. 长文本生成

八、文本生成应用场景

1. 内容创作
2. 对话系统
3. 数据增强
4. 代码生成

九、未来发展方向

1. 可控生成
2. 多模态生成
3. 高效训练
'''
    }
]

for doc_info in docs_content:
    doc = Document()
    
    heading = doc.add_heading(doc_info['title'], level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    paragraphs = doc_info['content'].split('\n\n')
    for para in paragraphs:
        p = doc.add_paragraph(para)
        p.style.font.size = Pt(11)
    
    doc.save(f'docs/{doc_info["filename"]}')
    print(f'已创建: docs/{doc_info["filename"]}')

print('\n所有DOCX文档创建完成！')